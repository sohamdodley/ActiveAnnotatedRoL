"""
ActiveReaderPresi
Assistive SQ3R + SPIDER Annotation Studio
Presidency University – Postgraduate Research Training

Core principle: The student performs every intellectual act.
The application only provides structure, storage, and visibility.

Pedagogical sources:
- Diana Ridley (2012), The Literature Review, Chapter 4
- SPIDER framework (Sample, Phenomenon of Interest, Design, Evaluation, Research type)
"""

import streamlit as st
import json
import io
from datetime import datetime

# Optional parsers
try:
    import fitz  # PyMuPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from ebooklib import epub
    from bs4 import BeautifulSoup
    HAS_EPUB = True
except ImportError:
    HAS_EPUB = False

# ------------------------------------------------------------
# Constants
# ------------------------------------------------------------
MAX_FILE_SIZE_MB = 25
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024

STAGES = ["Survey", "Question", "Read + Annotate", "Recall", "Review", "Write Review"]
SPIDER_KEYS = ["S", "P", "D", "E", "R"]
SPIDER_LABELS = {
    "S": "Sample",
    "P": "Phenomenon of Interest",
    "D": "Design",
    "E": "Evaluation",
    "R": "Research type",
}

# ------------------------------------------------------------
# Page config
# ------------------------------------------------------------
st.set_page_config(
    page_title="ActiveReaderPresi",
    page_icon="📖",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ------------------------------------------------------------
# Session state
# ------------------------------------------------------------
def init_state():
    defaults = {
        "mode": "student",
        "article_text": "",
        "article_title": "Untitled",
        "current_stage": "Survey",
        "survey_gist": "",
        "questions": ["", "", "", ""],
        "annotations": [],
        "recall": {k: "" for k in SPIDER_KEYS},
        "review_checks": {k: "" for k in SPIDER_KEYS},
        "final_review": "",
        "ai_flags": {stage: {"used": False, "note": ""} for stage in STAGES},
        "feedback": [],
        "completed_stages": [],
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()


def mark_complete(stage_name: str):
    cs = st.session_state.get("completed_stages", [])
    if not isinstance(cs, list):
        cs = list(cs)
    if stage_name not in cs:
        cs.append(stage_name)
    st.session_state.completed_stages = cs


def show_progress():
    current = st.session_state.current_stage
    try:
        idx = STAGES.index(current)
    except ValueError:
        idx = 0
    st.progress((idx + 1) / len(STAGES), text=f"Stage {idx + 1} of {len(STAGES)}: {current}")


# ------------------------------------------------------------
# DOCX Export Functions
# ------------------------------------------------------------
def create_docx_from_text(title: str, content: str) -> io.BytesIO:
    """Create a DOCX document from text content."""
    if not HAS_DOCX:
        return None
    
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    
    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_survey_docx() -> io.BytesIO:
    """Create DOCX for Survey stage."""
    if not HAS_DOCX:
        return None
    
    doc = Document()
    doc.add_heading(f"Survey: {st.session_state.get('article_title', 'Untitled')}", level=1)
    
    doc.add_heading("Article Gist", level=2)
    doc.add_paragraph(st.session_state.get("survey_gist", ""))
    
    doc.add_heading("AI Usage", level=2)
    ai_flag = st.session_state.ai_flags.get("Survey", {})
    doc.add_paragraph(f"AI Used: {'Yes' if ai_flag.get('used') else 'No'}")
    if ai_flag.get('note'):
        doc.add_paragraph(f"Details: {ai_flag.get('note')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_questions_docx() -> io.BytesIO:
    """Create DOCX for Question stage."""
    if not HAS_DOCX:
        return None
    
    doc = Document()
    doc.add_heading(f"Questions: {st.session_state.get('article_title', 'Untitled')}", level=1)
    
    doc.add_heading("Research Questions", level=2)
    for i, q in enumerate(st.session_state.get("questions", []), 1):
        if q.strip():
            doc.add_paragraph(q, style='List Number')
    
    doc.add_heading("AI Usage", level=2)
    ai_flag = st.session_state.ai_flags.get("Question", {})
    doc.add_paragraph(f"AI Used: {'Yes' if ai_flag.get('used') else 'No'}")
    if ai_flag.get('note'):
        doc.add_paragraph(f"Details: {ai_flag.get('note')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_annotations_docx() -> io.BytesIO:
    """Create DOCX for Read + Annotate stage."""
    if not HAS_DOCX:
        return None
    
    doc = Document()
    doc.add_heading(f"Annotations: {st.session_state.get('article_title', 'Untitled')}", level=1)
    
    doc.add_heading("Annotations", level=2)
    for i, ann in enumerate(st.session_state.get("annotations", []), 1):
        doc.add_heading(f"Annotation {i}", level=3)
        doc.add_paragraph(f"Category: {ann.get('color', 'Unknown')}")
        if ann.get('spider'):
            doc.add_paragraph(f"SPIDER Tag: {ann.get('spider')}")
        doc.add_paragraph(f"Text: {ann.get('text', '')}")
        if ann.get('comment'):
            doc.add_paragraph(f"Comment: {ann.get('comment')}")
    
    doc.add_heading("AI Usage", level=2)
    ai_flag = st.session_state.ai_flags.get("Read + Annotate", {})
    doc.add_paragraph(f"AI Used: {'Yes' if ai_flag.get('used') else 'No'}")
    if ai_flag.get('note'):
        doc.add_paragraph(f"Details: {ai_flag.get('note')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_recall_docx() -> io.BytesIO:
    """Create DOCX for Recall stage."""
    if not HAS_DOCX:
        return None
    
    doc = Document()
    doc.add_heading(f"Recall (SPIDER): {st.session_state.get('article_title', 'Untitled')}", level=1)
    
    for k in SPIDER_KEYS:
        doc.add_heading(f"{k} – {SPIDER_LABELS[k]}", level=2)
        doc.add_paragraph(st.session_state.recall.get(k, ""))
    
    doc.add_heading("AI Usage", level=2)
    ai_flag = st.session_state.ai_flags.get("Recall", {})
    doc.add_paragraph(f"AI Used: {'Yes' if ai_flag.get('used') else 'No'}")
    if ai_flag.get('note'):
        doc.add_paragraph(f"Details: {ai_flag.get('note')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_review_docx() -> io.BytesIO:
    """Create DOCX for Review stage."""
    if not HAS_DOCX:
        return None
    
    doc = Document()
    doc.add_heading(f"Review Verification: {st.session_state.get('article_title', 'Untitled')}", level=1)
    
    doc.add_heading("Verification Decisions", level=2)
    for k in SPIDER_KEYS:
        decision = st.session_state.review_checks.get(k, "")
        doc.add_paragraph(f"{k} – {SPIDER_LABELS[k]}: {decision if decision else '—'}")
    
    doc.add_heading("AI Usage", level=2)
    ai_flag = st.session_state.ai_flags.get("Review", {})
    doc.add_paragraph(f"AI Used: {'Yes' if ai_flag.get('used') else 'No'}")
    if ai_flag.get('note'):
        doc.add_paragraph(f"Details: {ai_flag.get('note')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_final_review_docx() -> io.BytesIO:
    """Create DOCX for Write Review stage."""
    if not HAS_DOCX:
        return None
    
    doc = Document()
    doc.add_heading(f"Review: {st.session_state.get('article_title', 'Untitled')}", level=1)
    
    doc.add_heading("Complete Review", level=2)
    doc.add_paragraph(st.session_state.get("final_review", ""))
    
    doc.add_heading("AI Usage", level=2)
    ai_flag = st.session_state.ai_flags.get("Write Review", {})
    doc.add_paragraph(f"AI Used: {'Yes' if ai_flag.get('used') else 'No'}")
    if ai_flag.get('note'):
        doc.add_paragraph(f"Details: {ai_flag.get('note')}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ------------------------------------------------------------
# Text extraction
# ------------------------------------------------------------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    if not HAS_PDF:
        return "[PDF support not available – install pymupdf]"
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = [page.get_text() for page in doc]
        doc.close()
        text = "\n\n".join(pages)
        if not text.strip():
            return "[PDF loaded but no text content extracted. File may be image-based or corrupted.]"
        return text
    except Exception as e:
        return f"[Error reading PDF: {str(e)}]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    if not HAS_DOCX:
        return "[DOCX support not available – install python-docx]"
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not text.strip():
            return "[DOCX loaded but no text content found in paragraphs.]"
        return text
    except Exception as e:
        return f"[Error reading DOCX: {str(e)}]"


def extract_text_from_epub(file_bytes: bytes) -> str:
    if not HAS_EPUB:
        return "[EPUB support not available – install ebooklib beautifulsoup4]"
    try:
        book = epub.read_epub(io.BytesIO(file_bytes))
        texts = []
        for item in book.get_items():
            if item.get_type() == 9:
                soup = BeautifulSoup(item.get_content(), "lxml")
                texts.append(soup.get_text(separator="\n", strip=True))
        text = "\n\n".join(texts)
        if not text.strip():
            return "[EPUB loaded but no text content extracted.]"
        return text
    except Exception as e:
        return f"[Error reading EPUB: {str(e)}]"


def extract_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    try:
        data = uploaded_file.read()
    except Exception as e:
        return f"[Error reading file: {str(e)}]"

    if len(data) > MAX_FILE_SIZE_BYTES:
        return f"File exceeds the {MAX_FILE_SIZE_MB} MB limit."

    if len(data) == 0:
        return "[File is empty.]"

    if name.endswith(".pdf"):
        return extract_text_from_pdf(data)
    elif name.endswith(".docx"):
        return extract_text_from_docx(data)
    elif name.endswith(".epub"):
        return extract_text_from_epub(data)
    elif name.endswith(".txt"):
        try:
            text = data.decode("utf-8", errors="replace")
            if not text.strip():
                return "[TXT file loaded but is empty.]"
            return text
        except Exception as e:
            return f"[Error reading TXT: {str(e)}]"
    else:
        return "[Unsupported format. Please upload PDF, DOCX, EPUB or TXT.]"


# ------------------------------------------------------------
# AI Usage Flag
# ------------------------------------------------------------
def render_ai_flag(stage: str):
    st.markdown("---")
    st.markdown(f"**AI Usage Flag — {stage}** (required)")
    used = st.radio(
        "Did you use AI in this stage?",
        ["No AI used", "AI used"],
        key=f"ai_radio_{stage}",
        horizontal=True,
        label_visibility="collapsed",
    )
    note = ""
    if used == "AI used":
        note = st.text_area(
            "Briefly describe how AI was used",
            key=f"ai_note_{stage}",
            height=70,
            placeholder="e.g. Checked a definition / Asked for grammar check of my own paragraph",
        )
    st.session_state.ai_flags[stage] = {
        "used": used == "AI used",
        "note": note if used == "AI used" else "",
    }


# ------------------------------------------------------------
# Sidebar (revised – more reliable on Streamlit Cloud)
# ------------------------------------------------------------
with st.sidebar:
    st.title("ActiveReaderPresi")
    st.caption("SQ3R + SPIDER · Assistive only")
    st.markdown(
        '<p style="font-size:0.72rem; color:#7A7A7A; line-height:1.35;">'
        "Based on Diana Ridley (2012), Ch. 4<br>SPIDER framework"
        "</p>",
        unsafe_allow_html=True,
    )
    st.markdown("---")

    # Stable mode selector
    mode_options = ["Student", "Supervisor"]
    current_mode = st.session_state.get("mode", "student").capitalize()
    if current_mode not in mode_options:
        current_mode = "Student"
    selected_mode = st.radio(
        "Mode",
        mode_options,
        index=mode_options.index(current_mode),
        key="mode_selector",
    )
    st.session_state.mode = selected_mode.lower()

    st.markdown("---")
    st.markdown("**Stages**")

    # Stage navigation – no explicit st.rerun() for reliability on Cloud
    for s in STAGES:
        is_done = s in st.session_state.get("completed_stages", [])
        label = f"✓ {s}" if is_done else s

        if s == st.session_state.get("current_stage"):
            st.markdown(f"**→ {label}**")
        else:
            if st.button(label, key=f"nav_{s}", use_container_width=True):
                st.session_state.current_stage = s

    st.markdown("---")

    if st.button("Reset all work", use_container_width=True):
        keep_mode = st.session_state.get("mode", "student")
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.session_state.mode = keep_mode
        st.session_state.current_stage = "Survey"
        st.rerun()

    if st.button("Export all work (JSON)", use_container_width=True):
        export = {
            "title": st.session_state.get("article_title", "Untitled"),
            "exported_at": datetime.utcnow().isoformat() + "Z",
            "survey_gist": st.session_state.get("survey_gist", ""),
            "questions": st.session_state.get("questions", []),
            "annotations": st.session_state.get("annotations", []),
            "recall": st.session_state.get("recall", {}),
            "review_checks": st.session_state.get("review_checks", {}),
            "final_review": st.session_state.get("final_review", ""),
            "ai_flags": st.session_state.get("ai_flags", {}),
            "feedback": st.session_state.get("feedback", []),
        }
        st.download_button(
            "Download JSON",
            data=json.dumps(export, indent=2, ensure_ascii=False),
            file_name="ActiveReaderPresi_export.json",
            mime="application/json",
            use_container_width=True,
        )


# ------------------------------------------------------------
# Supervisor view
# ------------------------------------------------------------
if st.session_state.mode == "supervisor":
    st.header("Supervisor Module")
    st.info("Read-only view. You may add formative feedback. You cannot edit the student’s intellectual content.")

    with st.expander("1. Survey", expanded=True):
        st.write(st.session_state.survey_gist or "_Not completed_")
        flag = st.session_state.ai_flags.get("Survey", {})
        st.caption(f"AI Flag: {'Yes — ' + flag.get('note', '') if flag.get('used') else 'No AI used'}")

    with st.expander("2. Question"):
        for i, q in enumerate(st.session_state.questions, 1):
            if q.strip():
                st.markdown(f"{i}. {q}")
        flag = st.session_state.ai_flags.get("Question", {})
        st.caption(f"AI Flag: {'Yes — ' + flag.get('note', '') if flag.get('used') else 'No AI used'}")

    with st.expander("3. Read + Annotate"):
        st.write(f"{len(st.session_state.annotations)} annotation(s)")
        for ann in st.session_state.annotations:
            st.markdown(
                f"- **[{ann.get('color', '')} · {ann.get('spider', '')}]** "
                f"{ann.get('comment') or ann.get('text', '')[:80]}"
            )
        flag = st.session_state.ai_flags.get("Read + Annotate", {})
        st.caption(f"AI Flag: {'Yes — ' + flag.get('note', '') if flag.get('used') else 'No AI used'}")

    with st.expander("4. Recall (SPIDER)"):
        for k in SPIDER_KEYS:
            st.markdown(f"**{k} – {SPIDER_LABELS[k]}**  \n{st.session_state.recall.get(k) or '_empty_'}")
        flag = st.session_state.ai_flags.get("Recall", {})
        st.caption(f"AI Flag: {'Yes — ' + flag.get('note', '') if flag.get('used') else 'No AI used'}")

    with st.expander("5. Review Decisions"):
        for k in SPIDER_KEYS:
            st.write(f"{k}: {st.session_state.review_checks.get(k) or '—'}")
        flag = st.session_state.ai_flags.get("Review", {})
        st.caption(f"AI Flag: {'Yes — ' + flag.get('note', '') if flag.get('used') else 'No AI used'}")

    with st.expander("6. Final Review"):
        st.write(st.session_state.final_review or "_Not written_")
        flag = st.session_state.ai_flags.get("Write Review", {})
        st.caption(f"AI Flag: {'Yes — ' + flag.get('note', '') if flag.get('used') else 'No AI used'}")

    st.markdown("---")
    st.subheader("Add Feedback")
    fb = st.text_area("Formative comment", height=100)
    if st.button("Save feedback"):
        if fb.strip():
            st.session_state.feedback.append({
                "ts": datetime.utcnow().isoformat() + "Z",
                "text": fb.strip(),
            })
            st.success("Feedback saved.")
            st.rerun()

    if st.session_state.feedback:
        st.markdown("**Previous feedback**")
        for f in st.session_state.feedback:
            st.markdown(f"- `{f['ts']}` — {f['text']}")

    st.stop()


# ------------------------------------------------------------
# Student – Upload
# ------------------------------------------------------------
if not st.session_state.article_text:
    st.header("ActiveReaderPresi")
    st.markdown("**Assistive SQ3R + SPIDER Annotation Studio**")
    st.caption("Core intellectual work remains the student’s. Maximum file size: 25 MB.")
    st.markdown(
        """
        <div style="font-size:0.85rem; color:#7A7A7A; margin-top:0.4rem;">
        Pedagogical sources: <strong>Diana Ridley (2012), <em>The Literature Review</em>, Chapter 4</strong>
        &nbsp;·&nbsp; <strong>SPIDER framework</strong>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.expander("Quick start – how this app works", expanded=True):
        st.markdown("""
1. **Upload** an article (PDF / DOCX / EPUB / TXT) or paste text.
2. Work through six stages: **Survey → Question → Read + Annotate → Recall → Review → Write Review**.
3. In **Read + Annotate** use the three colours recommended by Ridley to separate the author’s words from your own thinking.
4. In **Recall** the original text is hidden – write from memory under the five SPIDER headings.
5. Declare **AI usage** honestly at every stage.
6. Your supervisor can view all stages and add feedback.
7. Export your complete work as JSON when finished.

You can return to any stage at any time from the sidebar.
        """)

    col_upload, col_paste = st.columns(2)
    
    with col_upload:
        st.markdown("**Upload file**")
        uploaded = st.file_uploader(
            "Upload article or book chapter",
            type=["pdf", "docx", "epub", "txt"],
            help=f"Maximum size {MAX_FILE_SIZE_MB} MB",
            key="file_upload_widget",
        )

    with col_paste:
        st.markdown("**Or paste text**")
        paste = st.text_area("Or paste text here", height=200, key="paste_text_widget")

    # Process upload
    if uploaded is not None:
        with st.spinner("Extracting text…"):
            text = extract_text(uploaded)
            # Check if extraction returned an error message (starts with [)
            if text.startswith("["):
                st.error(f"❌ {text}")
            elif text.strip():
                st.session_state.article_text = text
                st.session_state.article_title = uploaded.name
                st.session_state.current_stage = "Survey"
                st.success(f"✅ Loaded: {uploaded.name} ({len(text):,} characters)")
                st.rerun()
            else:
                st.error("❌ Could not extract text from file. The file may be empty or corrupted.")

    # Process paste
    elif paste.strip() and st.button("Use pasted text", use_container_width=True, key="paste_btn"):
        st.session_state.article_text = paste.strip()
        st.session_state.article_title = "Pasted text"
        st.session_state.current_stage = "Survey"
        st.rerun()

    st.stop()


# ------------------------------------------------------------
# Student stages
# ------------------------------------------------------------
stage = st.session_state.current_stage
show_progress()
st.header(stage)

# ----- SURVEY -----
if stage == "Survey":
    st.markdown("""
    **Survey** (Diana Ridley, Chapter 4)  
    Skim the **entire** article in the left panel. Write your gist in the right panel at the same time.
    """)

    with st.expander("How to Survey (Ridley)", expanded=False):
        st.markdown("""
        - Look at the title and any abstract.
        - Scan all headings and sub-headings.
        - Read the opening paragraphs and the conclusion.
        - Notice the overall organisation.
        - Do **not** yet take detailed notes or annotate.
        - Write a short gist that captures what the article is about.
        """)

    # ---------- Parallel layout ----------
    st.markdown("**Split View: Skim left, write right. Keep both panels visible.**")
    left, right = st.columns([3, 2], gap="medium")

    with left:
        st.markdown("##### 📄 Article — skim here")
        # Use a disabled text_area for better scrolling reliability in Streamlit
        st.text_area(
            "Article text (scroll to read entire document)",
            value=st.session_state.article_text,
            height=520,
            disabled=True,
            label_visibility="collapsed",
        )
        st.caption(f"📊 {len(st.session_state.article_text):,} characters · Scroll freely to skim entire text")

    with right:
        st.markdown("##### ✍️ Your gist")
        gist_val = st.text_area(
            "Gist",
            value=st.session_state.survey_gist,
            height=420,
            placeholder="While skimming the article on the left, write what it is about in your own words…",
            label_visibility="collapsed",
            key="survey_gist_box",
        )
        st.session_state.survey_gist = gist_val
        
        # Word count tracker
        word_count = len(gist_val.split()) if gist_val.strip() else 0
        st.caption(f"💬 {word_count} words · Write here while you skim. Both panels stay visible.")

    # AI flag + save
    render_ai_flag("Survey")

    # Download button for Survey
    dl_col, save_col = st.columns(2)
    with dl_col:
        survey_docx = create_survey_docx()
        if survey_docx:
            st.download_button(
                "⬇️ Download as DOCX",
                data=survey_docx,
                file_name=f"Survey_{st.session_state.get('article_title', 'Untitled').replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    
    with save_col:
        if st.button("Save Survey & continue", type="primary", key="btn_save_survey", use_container_width=True):
            if not st.session_state.survey_gist.strip():
                st.warning("Please write a gist before continuing.")
            else:
                mark_complete("Survey")
                st.session_state.current_stage = "Question"
                st.rerun()

# ----- QUESTION -----
elif stage == "Question":
    st.markdown("Formulate questions you want the article to answer. The questions must be your own.")
    new_questions = []
    for i in range(max(4, len(st.session_state.questions))):
        q = st.text_input(
            f"Question {i+1}",
            value=st.session_state.questions[i] if i < len(st.session_state.questions) else "",
            key=f"q_{i}",
        )
        new_questions.append(q)
    st.session_state.questions = new_questions

    if st.button("Add another question"):
        st.session_state.questions.append("")
        st.rerun()

    render_ai_flag("Question")
    
    dl_col, save_col = st.columns(2)
    with dl_col:
        q_docx = create_questions_docx()
        if q_docx:
            st.download_button(
                "⬇️ Download as DOCX",
                data=q_docx,
                file_name=f"Questions_{st.session_state.get('article_title', 'Untitled').replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    
    with save_col:
        if st.button("Save Questions & continue", type="primary", use_container_width=True):
            mark_complete("Question")
            st.session_state.current_stage = "Read + Annotate"
            st.rerun()

# ----- READ + ANNOTATE -----
elif stage == "Read + Annotate":
    st.markdown("""
    **Active annotation** (following Diana Ridley, Chapter 4)  
    Separate the author’s words from your own thinking. Use colour to keep the distinction clear.
    """)
    st.markdown("""
    <div style="display:flex; gap:1.2rem; margin-bottom:0.8rem; font-size:0.88rem; flex-wrap:wrap;">
      <div><span style="background:#bbf7d0; padding:2px 8px; border-radius:4px;">Green – Source</span> Key claims, evidence, definitions</div>
      <div><span style="background:#bfdbfe; padding:2px 8px; border-radius:4px;">Blue – Structure</span> Organisation, section purpose, flow</div>
      <div><span style="background:#fecaca; padding:2px 8px; border-radius:4px;">Red – Response</span> Your critical comments, questions, links</div>
    </div>
    """, unsafe_allow_html=True)
    st.caption("Source: Diana Ridley (2012), Chapter 4. SPIDER tags may be added to any annotation.")

    st.text_area("Article text", value=st.session_state.article_text, height=280, disabled=True)

    with st.expander("Add new annotation", expanded=True):
        sel_text = st.text_area("Paste or type the passage you are annotating", height=70)
        color = st.selectbox(
            "Colour category (Ridley distinction)",
            [
                "Green – Source (author’s claims / evidence)",
                "Blue – Structure (organisation / flow)",
                "Red – Response (your critical comment)",
            ],
        )
        spider = st.selectbox(
            "SPIDER tag (optional)",
            ["", "S – Sample", "P – Phenomenon", "D – Design", "E – Evaluation", "R – Research type"],
        )
        comment = st.text_area("Your comment / note", height=70)
        if st.button("Add annotation"):
            if sel_text.strip():
                if color.startswith("Green"):
                    ckey = "Source"
                elif color.startswith("Blue"):
                    ckey = "Structure"
                else:
                    ckey = "Response"
                st.session_state.annotations.append({
                    "text": sel_text.strip(),
                    "color": ckey,
                    "spider": spider[:1] if spider else "",
                    "comment": comment.strip(),
                    "ts": datetime.utcnow().isoformat() + "Z",
                })
                st.success("Annotation added.")
                st.rerun()

    st.markdown(f"**Current annotations ({len(st.session_state.annotations)})**")
    for i, ann in enumerate(st.session_state.annotations):
        icon = {"Source": "🟢", "Structure": "🔵", "Response": "🔴"}.get(ann.get("color", ""), "⚪")
        st.markdown(
            f"{i+1}. {icon} `{ann['color']}` `{ann.get('spider') or '—'}` — "
            f"{ann.get('comment') or ann['text'][:100]}"
        )
        if st.button(f"Delete #{i+1}", key=f"del_ann_{i}"):
            st.session_state.annotations.pop(i)
            st.rerun()

    render_ai_flag("Read + Annotate")
    
    dl_col, save_col = st.columns(2)
    with dl_col:
        ann_docx = create_annotations_docx()
        if ann_docx:
            st.download_button(
                "⬇️ Download as DOCX",
                data=ann_docx,
                file_name=f"Annotations_{st.session_state.get('article_title', 'Untitled').replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    
    with save_col:
        if st.button("Save Annotations & continue", type="primary", use_container_width=True):
            mark_complete("Read + Annotate")
            st.session_state.current_stage = "Recall"
            st.rerun()

# ----- RECALL -----
elif stage == "Recall":
    st.warning("Source text is hidden. Write from memory only under the five SPIDER headings.")
    with st.expander("How to use this stage (Ridley + SPIDER)"):
        st.markdown("""
- Close the original text in your mind.
- Write what you remember under each SPIDER heading.
- Do **not** look back at the article or your annotations yet.
- After you finish, move to Review to check accuracy.
- This stage trains active recall (Ridley, Chapter 4).
        """)
    for k in SPIDER_KEYS:
        st.session_state.recall[k] = st.text_area(
            f"{k} – {SPIDER_LABELS[k]}",
            value=st.session_state.recall.get(k, ""),
            height=90,
            key=f"recall_{k}",
        )
    render_ai_flag("Recall")
    
    dl_col, save_col = st.columns(2)
    with dl_col:
        recall_docx = create_recall_docx()
        if recall_docx:
            st.download_button(
                "⬇️ Download as DOCX",
                data=recall_docx,
                file_name=f"Recall_{st.session_state.get('article_title', 'Untitled').replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    
    with save_col:
        if st.button("Save Recall & continue", type="primary", use_container_width=True):
            mark_complete("Recall")
            st.session_state.current_stage = "Review"
            st.rerun()

# ----- REVIEW -----
elif stage == "Review":
    st.markdown("Compare your Recall notes with the original text. Mark each SPIDER element.")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Your Recall**")
        for k in SPIDER_KEYS:
            st.markdown(f"**{k}**  \n{st.session_state.recall.get(k) or '_empty_'}")
    with col2:
        st.markdown("**Original (excerpt)**")
        st.text_area(
            "Source",
            value=st.session_state.article_text[:4000],
            height=380,
            disabled=True,
            label_visibility="collapsed",
        )

    st.markdown("**Verification decisions**")
    for k in SPIDER_KEYS:
        st.session_state.review_checks[k] = st.radio(
            f"{k} – {SPIDER_LABELS[k]}",
            ["", "Accurate", "Incomplete", "Needs revision"],
            key=f"check_{k}",
            horizontal=True,
        )

    render_ai_flag("Review")
    
    dl_col, save_col = st.columns(2)
    with dl_col:
        review_docx = create_review_docx()
        if review_docx:
            st.download_button(
                "⬇️ Download as DOCX",
                data=review_docx,
                file_name=f"Review_{st.session_state.get('article_title', 'Untitled').replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    
    with save_col:
        if st.button("Save Review & continue", type="primary", use_container_width=True):
            mark_complete("Review")
            st.session_state.current_stage = "Write Review"
            st.rerun()

# ----- WRITE REVIEW -----
elif stage == "Write Review":
    st.markdown("Write the complete review using only your own previous work. The app supplies structure; you write the prose.")
    with st.expander("Your accumulated notes", expanded=False):
        st.markdown(f"**Gist**  \n{st.session_state.survey_gist or '—'}")
        st.markdown("**Questions**")
        for q in st.session_state.questions:
            if q.strip():
                st.markdown(f"- {q}")
        st.markdown("**SPIDER Recall**")
        for k in SPIDER_KEYS:
            st.markdown(f"- **{k}**: {st.session_state.recall.get(k) or '—'}")
        st.markdown(f"**Annotations**: {len(st.session_state.annotations)} item(s)")

    st.session_state.final_review = st.text_area(
        "Final review",
        value=st.session_state.final_review,
        height=380,
        placeholder="Write your complete review of the article here…",
    )
    render_ai_flag("Write Review")
    
    # Download button
    final_docx = create_final_review_docx()
    if final_docx:
        st.download_button(
            "⬇️ Download Final Review as DOCX",
            data=final_docx,
            file_name=f"FinalReview_{st.session_state.get('article_title', 'Untitled').replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    
    col_a, col_b = st.columns(2)
    with col_a:
        if st.button("Save draft", use_container_width=True):
            mark_complete("Write Review")
            st.success("Draft saved.")
    with col_b:
        if st.button("Mark complete", type="primary", use_container_width=True):
            mark_complete("Write Review")
            st.success("Review marked complete. You may still return to any stage and iterate. Use the sidebar to Export.")
